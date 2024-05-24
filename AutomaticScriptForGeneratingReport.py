import pandas as pd
import os
import matplotlib
matplotlib.use('Agg')
import matplotlib.pyplot as plt 
 
import sys
import numpy as np
import docx 
from docx.shared import Cm, RGBColor,Pt
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml
from datetime import datetime
from docx.enum.text import WD_ALIGN_PARAGRAPH
from tkinter import *  
from tkinter import filedialog
from tkinter import messagebox
from threading import Thread

PresentMonDataDir='PresentMon'
XAxisNum = 100
TestPlatformParameters = ['OS','CPU Type','CPU NumberOfCores','CPU NumberOfLogicalProcessors','Baseboard','SMBIOSBIOSVersion','RAM Capacity','RAM ConfiguredClockSpeed','RAM Manufacturer','GPU','DriverVersion']
TestCaseParameters = ['Average GPU duration','Average Frametime','Aveage FPS','Ratio over 30 FPS','Dropped %']
test_case_list = []
test_folder_list = []
root_path=''
## 
# this class is to store the data of each test case        
class TestCase :
    def __init__(self,case_name) :
        self.name = case_name
        self.round_number = 0
        self.round_data_tuple_list = []
        self.rounds_diagram_path = ''
        self.median_round_frametime_chart_path=''
        self.median_round_presentmon_data_path=''


##
# this class is to map to the structure of test report doc and store all related data. 
class ReportContent:
    def __init__(self,name) :
        self.title = name
        self.test_information = 'Test Information'
        self.test_date = ''
        self.test_platform_list=[]
        self.test_setting_path_list=[]
        self.test_scene_path_list = []
        self.test_case_list=[]
        self.data_compare_tuple_list=[]
        self.emon_data='Emon Data'
        self.Xperf_data='Xperf Data'
    def clear(self):
        self.title = ' '
        self.test_information = 'Test Information'
        self.test_date = ''
        self.test_platform_list=[]
        self.test_setting_path_list=[]
        self.test_scene_path_list = []
        self.test_case_list=[]
        self.data_compare_tuple_list=[]
        self.emon_data='Emon Data'
        self.Xperf_data='Xperf Data'
 

def parse_config_file(file_path) :
    flag = ''
    f= open(file_path,'r',encoding='utf-8' )  
    keyword_list = ['[OS]','[CPU]','[Baseboard]','[BIOS]','[RAM]','[GPU]']
    info_list = []    
    temp_list =[]
    flag = False
    while True:    
        linedata = f.readline() 
        if flag:
            if not linedata.strip() == '' :
                temp_list.append(linedata.strip()) 

        if linedata.strip() in keyword_list: 
            temp_list =[]
            flag = True

        if linedata.strip() == ''and flag:
            flag = False
            info_list.append(temp_list)          
        if not linedata :
            break
    f.close() 
    
    platform_info_list = []
    for index, para_list in enumerate(info_list): 
        # get windows version
        if(index == 0) :
            
            str1= ''
            str2 =''
            for val in para_list : 
                if 'Name' in val :
                    str1 = val.split('=')[1] 
                if 'Build' in val:
                    str2 = val.split('=')[1]
            platform_info_list.append(str1+' '+str2) 
        # get CPU info
        if(index == 1) :
            
            str1= ''
            str2 =''
            str3=''
            for val in para_list :
                if 'Name' in val :
                    str1 = val.split('=')[1]
                if 'NumberOfCores' in val:
                    str2 = val.split('=')[1]
                if 'NumberOfLogicalProcessors' in val:
                    str3 = val.split('=')[1]
            platform_info_list.append(str1)                
            platform_info_list.append(str2)                
            platform_info_list.append(str3) 
        # get baseboard info
        if(index == 2) :
            str1= '' 
            for val in para_list :
                if 'Product' in val :
                    str1 = val.split('=')[1] 
            platform_info_list.append(str1) 
        # get bios info
        if(index == 3) :
            str1= '' 
            for val in para_list :
                if 'SMBIOSBIOSVersion' in val :
                    str1 = val.split('=')[1] 
            platform_info_list.append(str1)
        # get ram info
        if(index == 4) :
            
            str1= ''
            str2 =''
            str3=''
            for val in para_list :
                if 'Capacity' in val :
                    str1 = str( round(float(val.split('=')[1])/(1024*1024*1024))) + 'GB'
                if 'ConfiguredClockSpeed' in val:
                    str2 = val.split('=')[1]
                if 'Manufacturer' in val:
                    str3 = val.split('=')[1]
            platform_info_list.append(str1)                
            platform_info_list.append(str2)                
            platform_info_list.append(str3)
        # get GPU info
        if(index == 5) :
            
            str1= ''
            str2 =''
            for val in para_list :
                if 'Name' in val :
                    str1 = val.split('=')[1]
                if 'DriverVersion' in val:
                    str2 = val.split('=')[1]
            platform_info_list.append(str1)
            platform_info_list.append(str2)
            
    # save to doc data structure.
    total_doc_content.test_platform_list.append(platform_info_list)

def get_platforms_list():

    lnk_files_list = sorted([ name for name in os.listdir(root_path) if name.endswith('.lnk')])  
    
    global test_case_list 
    global test_folder_list

    
    test_case_list = []
    test_folder_list = []
    temp_folder_list =  sorted([ dir for dir in os.listdir(root_path) if os.path.isdir(os.path.join(root_path,dir))]) 
    
    #filter invaild test case 
    for dir in temp_folder_list :
        for case in lnk_files_list :
            case=case.strip('.lnk')  
            templist = case.split('---') 
            if os.path.exists(os.path.join(root_path,dir,PresentMonDataDir)) and dir in case:
                test_case_title = templist[3]+' '+templist[5]+' '+templist[8]+' '+templist[6]+' '+templist[7]+' '+templist[9]+' '+templist[-1]
                test_case_list.append(test_case_title)
                test_folder_list.append(dir)
                break

def extract_case_information(selection):
    files = sorted([ name for name in os.listdir(root_path) if name.endswith('.lnk')])
    dir_list =[]
    test_file_list =[]
    test_case_title= ''
    
    dir_list=[test_folder_list[i] for i in selection] 
    
    textvr_global.set('Extract test cases information')
    for dir in dir_list :
        
        dir_path = os.path.join(root_path,dir,PresentMonDataDir) 
        if(os.path.exists(dir_path)):
            for file in files:
                if dir in file :  
                    file =  file.strip('.lnk')                 
                    templist = file.split('---')
                    test_case_title = templist[3]+' '+templist[5]+' '+templist[8]+' '+templist[6]+' '+templist[7]+' '+templist[9]+' '+templist[-1]

                    #get title name
                    total_doc_content.title = templist[3]
                    #get test date
                    total_doc_content.test_date = templist[4].split('-')[0]
                    break
            
            #append presentmon data file to file list
            for path in os.listdir(dir_path):
                # check if current path is a file
                if os.path.isfile(os.path.join(dir_path, path)):
                    test_file_list.append(os.path.join(dir_path, path))   
         
        parse_presentdata_files(test_file_list,test_case_title)
        test_file_list = []
        #parse system configuration informaion
        config_file_path = os.path.join(root_path,dir,'SystemReport-01.txt')
        if(os.path.exists(config_file_path)):
            parse_config_file(config_file_path)
        
        if len(total_doc_content.test_setting_path_list)<=0 :
            dir_path = os.path.join(root_path,dir ) 
            img_path_list = sorted([os.path.join(dir_path,name) for name in os.listdir(dir_path) if name.endswith('.png')or name.endswith('.jpg')]) 
            for img in img_path_list :
            # get test setting screenshot picture file
                if 'setting' in img.lower() :
                    total_doc_content.test_setting_path_list.append(img)
            # get test scene screenshot picture file
                if 'screenshot' in img.lower():
                    total_doc_content.test_scene_path_list.append(img)
    #return testfiles


def parse_presentdata_files(files,title):

    
    textvr_global.set('Parsing '+title)
    file_list_len=len(files)
    
    testcase = TestCase(title)
    testcase.round_number=file_list_len
 
    #testcase.round_data_tuple_list
    if(file_list_len==0):
        return
    else:        
        plt.figure(dpi=200,figsize=(16,8))
        for index, file in enumerate( files ): 
         
            df=pd.read_csv(file)
            stride = int( len(df.msGPUActive)/XAxisNum)

            round_data_tuple  = []
            round_data_tuple.append("%.2f" %(df.msGPUActive.mean()))
            round_data_tuple.append("%.2f" %(df.msBetweenPresents.mean()))
            round_data_tuple.append("%.2f" %(1000/float(df.msBetweenPresents.mean())))
            #compute ratio of fps > 30fps
            count = len([x for x in df.msBetweenPresents.tolist() if x <= 1000/30])
            round_data_tuple.append("%.1f%%" % (count/len(df.msBetweenPresents) * 100))            
            round_data_tuple.append("%.2f" %(df.Dropped.sum()))
            testcase.round_data_tuple_list.append(round_data_tuple)

            # draw all round diagram for one test case 
            x = []
            GPUDuration = []
            GPUDurationAverage = []
            MsBetweenPresents= []
            MsBetweenPresentsAverage= []
            for i in range(XAxisNum):
                x.append(i)
                GPUDuration.append(df.msGPUActive[i*stride])
                GPUDurationAverage.append(round(df.msGPUActive.mean(),2))         

                MsBetweenPresents.append(df.msBetweenPresents[i*stride])
                MsBetweenPresentsAverage.append(round(df.msBetweenPresents.mean(),2))   
            ax=plt.subplot(2,int(file_list_len/2)+1,index+1)
            ax.plot(x,GPUDuration,label= "GPUDuration" )
            ax.plot(x,GPUDurationAverage,'b' ,label= "GPUDurationAverage")
            ax.plot(x,MsBetweenPresents,'orange' ,label= "MsBetweenPresents")
            ax.plot(x,MsBetweenPresentsAverage,'brown',label= "MsBetweenPresentsAverage" )
            ax.set_ylim(bottom=0.)
            ax.set_title(file.strip('.csv').split('\\')[-1])
            ax.set_xlabel("Frame No.")
            ax.set_ylabel("FrameTime(ms)")
            ax.legend( loc='lower right',fontsize="x-small") 
        plt.suptitle(title)
        #plt.show()
        fname=title+'.png'
        plt.savefig(fname)
        plt.close()
        testcase.rounds_diagram_path=fname
    # get median round presentmon data file path
    average_fps_list = []
    for case in testcase.round_data_tuple_list :
        average_fps_list.append(case[2]) 
    sorted_average_fps_list = sorted(average_fps_list)
    testcase.median_round_presentmon_data_path=files[average_fps_list.index(sorted_average_fps_list[int(len(sorted_average_fps_list)/2)])]
    #
    total_doc_content.test_case_list.append(testcase)

# doc heading 3 style. 
def add_heading3(doc, text) :
    table = doc.add_table(rows=1,cols=1,style='Light List Accent 5')
    cell = table.cell(0,0)    
    #shading_elm_1 = parse_xml(r'<w:shd {} w:fill="98F5FF"/>'.format(nsdecls('w')))
    #cell._tc.get_or_add_tcPr().append(shading_elm_1)
    cell.text=text
    cell.paragraphs[0].runs[0].font.size = Pt(16)   
    #cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)  
    doc.add_paragraph(" ")
#
def get_platform_name(testcase):
    keyword_list = ['XeSS','DLSS','8+0','8+8','8+16']
    trim_keyword_list = ['Nvidia_','RX_','AMD_','_C1_rev_8','_IBC']

    for kwd in trim_keyword_list :
        testcase=testcase.replace(kwd,'')
    
    temp = testcase.split(' ')
    name = temp[1]+ ' '+temp[2]

    for kwd in keyword_list:
        if kwd in testcase:
            name = name+' '+kwd  
    return name
def save_test_summary_bar(doc) :
    testcase_bar_data_list = []    
    bar_percenttile=[1,90]
    bar_x = np.arange(len(total_doc_content.test_case_list))
    bar_x_labels =[]
    bar_average_fps=[]
    for i,testcase in enumerate(total_doc_content.test_case_list) :
        #doc.add_paragraph(testcase.median_round_presentmon_data_path)
        df = pd.read_csv(testcase.median_round_presentmon_data_path)
        frametime_list = df.msBetweenPresents.to_list() 
        fps_list = [1000/i for i in frametime_list] 
        data = np.array(fps_list)
        result = np.percentile(data,bar_percenttile)
        bar_average_fps.append(round(1000/np.array(frametime_list).mean(),2))
        bar_x_labels.append(get_platform_name(testcase.name))#'Platform%d'%(i+1))#testcase.name)
        testcase_bar_data_list.append(result)
    # begin to draw bar chart 
    bar_total_width = 1
    bar_each_width = bar_total_width/(len(bar_percenttile)+3)
    plt.figure(dpi=200,figsize=(16,8))
    plt.bar(bar_x,bar_average_fps,width=bar_each_width,label='Average FPS')
    for i  in range(len(bar_x)) :
        plt.text(bar_x[i], bar_average_fps[i]+2,bar_average_fps[i],ha='center',fontsize=14)
    for i in range(len(bar_percenttile)) :
        percentile_fps_list = [p[i] for p in testcase_bar_data_list]
        plt.bar(bar_x+(i+1)*bar_each_width,percentile_fps_list,width=bar_each_width, label='%d'%bar_percenttile[i]+'% FPS')
        for j in range(len(bar_x)):
            plt.text(bar_x[j]+(i+1)*bar_each_width, percentile_fps_list[j]+2,round(percentile_fps_list[j],2),ha='center',fontsize=14)
    plt.xticks(bar_x,bar_x_labels,fontsize=14)
    plt.legend()
    plt.ylabel("FPS",fontsize=20) 
    plt.savefig('bar_temp.png')
    doc.add_picture('bar_temp.png' ,width=Cm(15))
    os.remove('bar_temp.png')


#  save each case informaiton to doc 
def save_each_case_detail_analysis_report(doc) :
    for i,testcase in enumerate(total_doc_content.test_case_list) :     

        textvr_global.set('Store all rounds informations of '+testcase.name)   
        doc.add_heading(testcase.name,level=3)
        doc.add_heading('Data statistic',level=4)
        doc.add_paragraph('In this case, we have run %d times and conclude this data as below.'%testcase.round_number)
        testcase_table=doc.add_table(rows=testcase.round_number+1,cols=len(TestCaseParameters)+1)
        testcase_table.style = 'Medium Grid 3 Accent 1'
        # fill out first row
        for j,para in enumerate( TestCaseParameters):
            testcase_table.rows[0].cells[j+1].text = para
        # fill out the other row
        for k, data in enumerate(testcase.round_data_tuple_list) :            
            testcase_table.rows[k+1].cells[0].text = 'Round%d'%(k+1)
            for m,value in enumerate(data):
                cell=testcase_table.rows[k+1].cells[m+1]
                cell.text =value
                cell.paragraphs[0].runs[0].font.size = Pt(9)  
        
        doc.add_heading('Diagram all rounds',level=4)
        doc.add_picture(testcase.rounds_diagram_path,width=Cm(15))
        os.remove(testcase.rounds_diagram_path)
        
        textvr_global.set('Store median round informations of '+testcase.name)  
        #add each case median round frametime sensitivity analysis
        doc.add_heading('Median Round Frametime Sensitivity Analysis',level=4)
        median_frametime_table = doc.add_table(rows=14, cols=2,style = "Light Grid Accent 6")
        median_frametime_table.columns[0].cells[0].text = "Total Frames"
        median_frametime_table.columns[0].cells[1].text = "Total Duration"
        median_frametime_table.columns[0].cells[2].text = "Avg. GPU Utilization"
        median_frametime_table.columns[0].cells[3].text = "Avg. Frametime"
        median_frametime_table.columns[0].cells[4].text = "Avg. FPS"
        median_frametime_table.columns[0].cells[5].text = "Frames Dropped"
        median_frametime_table.columns[0].cells[6].text = "Frames GPU Bound Number"
        median_frametime_table.columns[0].cells[7].text = "Frames GPU Bound Percent"
        median_frametime_table.columns[0].cells[8].text = "Mean GPU Time"
        median_frametime_table.columns[0].cells[9].text = "Min GPU Time"        
        median_frametime_table.columns[0].cells[10].text = "Max GPU Time"        
        median_frametime_table.columns[0].cells[11].text = "Mean CPU Time"
        median_frametime_table.columns[0].cells[12].text = "Min CPU Time"        
        median_frametime_table.columns[0].cells[13].text = "Max CPU Time"

        for i in range(0,14):
            median_frametime_table.columns[0].cells[i].paragraphs[0].alignment = 2 

        df = pd.read_csv(testcase.median_round_presentmon_data_path)
        frametime_list= df.msBetweenPresents.to_list()
        gpuduration_list = df.msGPUActive.to_list()
        TimeInSeconds_list = df.TimeInSeconds.to_list()
        Dropped_list = df.Dropped.to_list() 

        median_frametime_table.columns[1].cells[0].text = str(len(frametime_list))
        median_frametime_table.columns[1].cells[1].text = str(round(TimeInSeconds_list[-1])) +' seconds'
        median_frametime_table.columns[1].cells[2].text = str(round(100*np.array(gpuduration_list).mean()/np.array(frametime_list).mean(),2))+'%'
        median_frametime_table.columns[1].cells[3].text =str(round(np.array(frametime_list).mean(),2))+' ms'
        median_frametime_table.columns[1].cells[4].text =str(round(1000/np.array(frametime_list).mean(),2))
        median_frametime_table.columns[1].cells[5].text =str(np.array(Dropped_list).sum())

        bound_num=0
        bound_list = []
        bound_percent_list = []
        for i in range(len(frametime_list)):
            bound_percent_list.append(min(100, 100*gpuduration_list[i]/frametime_list[i]))         
            if gpuduration_list[i]/frametime_list[i]>=0.95 :
                bound_num=bound_num+1
                bound_list.append(1)
            else:
                bound_list.append(0)

        median_frametime_table.columns[1].cells[6].text =str(bound_num)
        median_frametime_table.columns[1].cells[7].text =str(round(100*bound_num/len(frametime_list),2))+'%'
        median_frametime_table.columns[1].cells[8].text = str(round(np.array(gpuduration_list).mean(),2))+' ms'
        median_frametime_table.columns[1].cells[9].text =str(round(np.array(gpuduration_list).min(),2))+' ms'
        median_frametime_table.columns[1].cells[10].text =str(round(np.array(gpuduration_list).max(),2))+' ms'
        median_frametime_table.columns[1].cells[11].text =str(round(np.array(frametime_list).mean(),2))+' ms'
        median_frametime_table.columns[1].cells[12].text =str(round(np.array(frametime_list).min(),2))+' ms'
        median_frametime_table.columns[1].cells[13].text =str(round(np.array(frametime_list).max(),2))+' ms'

        textvr_global.set('Store median round frametime diagram of '+testcase.name)  
        plt.figure(dpi=100,figsize=(24,8)) 
        plt.plot(TimeInSeconds_list[0:-1:50],frametime_list[0:-1:50],label = 'CPU Frametime' )
        plt.plot(TimeInSeconds_list[0:-1:50],gpuduration_list[0:-1:50],label = 'GPU Frametime' )
        plt.xlabel('Time In Seconds')
        plt.ylabel('Milliseconds')
        plt.legend()
        plt.savefig('median_frametime_compare.png')
        doc.add_picture('median_frametime_compare.png' ,width=Cm(15))
        os.remove('median_frametime_compare.png')
        plt.close()
#gpu bound
        textvr_global.set('Store median round GPU bound diagram of '+testcase.name)  
        plt.figure(dpi=100,figsize=(16,2)) 
        plt.bar(TimeInSeconds_list,bound_list,width=0.1,label = 'GPU Bound' ) 
        plt.xlabel('Time In Seconds')
        plt.ylabel('GPU BOUND')
        plt.legend(loc=1)
        plt.savefig('temp_gpu_bound.png')
        doc.add_picture('temp_gpu_bound.png' ,width=Cm(15))
        os.remove('temp_gpu_bound.png')
        plt.close()
# gpu utilization
        textvr_global.set('Store median round GPU utilization diagram of '+testcase.name)  
        plt.figure(dpi=100,figsize=(16,2)) 
        plt.bar(TimeInSeconds_list,bound_percent_list,width=0.1,label = 'GPU Utilization %' ) 
        plt.xlabel('Time In Seconds')
        plt.ylabel('GPU UTILIZATION')
        plt.legend(loc=1)
        plt.savefig('temp_gpu_utilizaiton.png')
        doc.add_picture('temp_gpu_utilizaiton.png' ,width=Cm(15))
        os.remove('temp_gpu_utilizaiton.png')
        plt.close()
# CPU Time / GPU Time Correlation Coefficient
        textvr_global.set('Store median round CPU Time/GPU Time Correlation Coefficient of '+testcase.name)  
        plt.figure(dpi=100,figsize=(16,8)) 
        
        cpu_smart_target = get_smart_target(frametime_list)
        gpu_smart_target = get_smart_target(gpuduration_list)
 
        max_cpu_frametime = max(frametime_list)
        max_gpu_frametime = max(gpuduration_list)
 
        
        plt.plot(np.linspace(0,max_gpu_frametime,2),np.repeat(cpu_smart_target,2),color='red',label='Samrt Target')
        plt.plot(np.repeat(gpu_smart_target,2),np.linspace(0,max_cpu_frametime,2),color='red') 

        plt.fill_between(np.linspace(0,gpu_smart_target,2),0,np.repeat(cpu_smart_target,2),color='green',label='Good Perf.')
        plt.fill_between(np.linspace(0,gpu_smart_target,2),cpu_smart_target,np.repeat(max_cpu_frametime,2),color='yellow',label='Low CPU Perf.')
        plt.fill_between(np.linspace(gpu_smart_target,max_gpu_frametime,2),0,np.repeat(cpu_smart_target,2),color='orange',label='Low GPU Perf.')
        plt.fill_between(np.linspace(gpu_smart_target,max_gpu_frametime,2),cpu_smart_target,np.repeat(max_cpu_frametime,2),color='pink',label='Low Perf.')

        plt.scatter(gpuduration_list ,frametime_list ,s=10,color='blue',label='Frame' ) 

        plt.xlabel('GPU Time (ms)')
        plt.ylabel('CPU Time (ms)')
        plt.legend()
        plt.savefig('gpu_cpu_co.png')
        doc.add_picture('gpu_cpu_co.png' ,width=Cm(15))
        os.remove('gpu_cpu_co.png')
        plt.close() 
# summary test platforms infromation
def get_smart_target(frametime_list):
    average_frametime = np.array(frametime_list).mean() 
    smart_fps = 990
    while 1:
        if 1000/smart_fps >= average_frametime :
            break
        else:
            if smart_fps < 60 :
                break
            smart_fps = smart_fps -30 
    return 1000/smart_fps
        
def save_platforms_information(doc):

#add platform table.
    platform_table=doc.add_table(rows=len(TestPlatformParameters)+1, cols= len(total_doc_content.test_platform_list)+1)#,style="Table Grid")
    platform_table.style = 'Medium Grid 3 Accent 1'
#set first column content    
    for index, para in enumerate(TestPlatformParameters ):
        platform_table.columns[0].cells[index+1].text = para

# set other column content of platform table. 
    for index, platform in  enumerate(total_doc_content.test_platform_list) :
        platform_table.rows[0].cells[index+1].text = 'Platform'+str(index+1)
        for i,val in enumerate(platform):
            cell=platform_table.columns[index+1].cells[i+1]
            cell.text = val
            cell.paragraphs[0].runs[0].font.size = Pt(9)   

#
def save_test_config_table(doc):
# set config parameter table(2*4) 
    config_table = doc.add_table(rows=4, cols=2,style = "Light Grid Accent 6")
    config_table.columns[0].cells[0].text = "QUALITY LEVEL"
    config_table.columns[0].cells[1].text = "API VERSION"
    config_table.columns[0].cells[2].text = "RESOLUTION"
    config_table.columns[0].cells[3].text = "GOAL"
    for i in range(0,4):
        config_table.columns[0].cells[i].paragraphs[0].alignment = 2 
    if len(total_doc_content.test_case_list) >0 :        
        
        temp_config_list = total_doc_content.test_case_list[0].name.split(' ')
        config_table.columns[1].cells[0].text =  temp_config_list[4]
        config_table.columns[1].cells[1].text =  temp_config_list[5]
        config_table.columns[1].cells[2].text =  temp_config_list[3]
        config_table.columns[1].cells[3].text =  temp_config_list[1]

    for test_setting_path in total_doc_content.test_setting_path_list:
        doc.add_picture(test_setting_path, width=Cm(15))
#
def save_all_cases_frametime_diagram(doc):
# frametime chart of all cases
    plt.figure(dpi=100,figsize=(16,8))
    for i,testcase in enumerate(total_doc_content.test_case_list) :
        df = pd.read_csv(testcase.median_round_presentmon_data_path)
        x = df.TimeInSeconds.to_list()
        y = df.msBetweenPresents.to_list()
        plt.plot(x[0:-1:50],y[0:-1:50],label = testcase.name )#linewidth=0.2,
    plt.xlabel('Time In Seconds')
    plt.ylabel('Milliseconds')
    plt.legend()
    plt.savefig('frametime_compare.png')
    doc.add_picture('frametime_compare.png' ,width=Cm(15))
    os.remove('frametime_compare.png')
    plt.close()


# generate word file
def save_word_file():
    doc = docx.Document()
    textvr_global.set('Store heading line ')
#set heading 1 style
    heading_1_style =  doc.styles['Heading 1']
    heading_1_style.font.color.rgb = RGBColor(0, 0, 0)
    heading_1_style.font.size = Pt(24)
    heading_1_style.font.name = 'Calibri Light'
    heading_1_style.paragraph_format.alignment = 1
    doc.add_heading((total_doc_content.title+' Performance testing Report').upper(),level =1)
    
    doc.add_paragraph( datetime.now().strftime('%Y-%m-%d')).alignment = WD_ALIGN_PARAGRAPH.CENTER
    #doc.add_paragraph( "DO NOT COPY OR DISTRIBUTE ")
    #doc.add_paragraph( 'Intel and Tencent Confidential ')
    paragraph = doc.add_paragraph()
    sentence = paragraph.add_run('DO NOT COPY OR DISTRIBUTE \n Intel Confidential')
    sentence.italic = True
    sentence.font.color.rgb = RGBColor(255, 0, 0)
    sentence.font.name = 'Calibri Light'
    sentence.font.size= Pt(14)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_page_break()

    doc.add_heading(total_doc_content.test_information,level = 2)

    doc.add_heading('Test Date', level=3)
    doc.add_paragraph('This case was test on '+total_doc_content.test_date)
    

    #doc.add_heading('SYSTEM CONFIGURATIONS', level=3)
    textvr_global.set('Store system configurations ')
    add_heading3(doc,'SYSTEM CONFIGURATIONS' )
    save_platforms_information(doc)

# set test setting screenshot
    #doc.add_heading('GAME SETTING', level=3)
    textvr_global.set('Store game setting')
    doc.add_paragraph(" ")
    add_heading3(doc, 'GAME SETTING')
    save_test_config_table(doc)

# set test scene screenshot

    textvr_global.set('Store game screenshot')
    doc.add_heading('GAME SCENE SCREENSHOT', level=3)
    for test_scene_path in total_doc_content.test_scene_path_list:
        doc.add_picture(test_scene_path, width=Cm(15))
# set test summary    
    textvr_global.set('Store test summary')
    add_heading3(doc, 'TESTING RESULTS')
    doc.add_paragraph('SUMMARY') 
    #collect all data needed to plot bar 
    save_test_summary_bar(doc)
    save_all_cases_frametime_diagram(doc)

# visualize presentmon data
    doc.add_heading('Detail performance data based on PresentMon',level = 2)
    save_each_case_detail_analysis_report(doc)
# Emon data analysis
    doc.add_heading('Emon data analysis',level = 2)
    doc.add_paragraph('N/A')
# GPA analysis
    doc.add_heading('GPA analysis',level = 2)
    doc.add_paragraph('N/A')
# ETL analysis
    doc.add_heading('ETL analysis',level = 2)
    doc.add_paragraph('N/A')

    filename = total_doc_content.title+'.docx'
    if(os.path.exists(filename)) :
        os.remove(filename)
    doc.save(filename)

def button_choose_dir_fun():    
    
    listbox_global.delete(0, END)
    global root_path
    root_path=filedialog.askdirectory() 
    #label_global.config(text=root_path)
    if root_path !='' :
        get_platforms_list()

        if len(test_case_list)>0 :
            for item in test_case_list:
                listbox_global.insert(END, item)
        else:
            messagebox.showinfo("Warning", "Please choose right folder")
def thread_fun():

    selection = listbox_global.curselection() 
    if len(selection) ==0 :
        selection =[i for i in range(0,listbox_global.size())]  
    extract_case_information(selection)    
    save_word_file() 
    listbox_global.delete(0, END)
    textvr_global.set('Finish!')
    # empty all content.
    total_doc_content.clear()
    
    button_choose_dir.config(state='normal')
    button_process.config(state='normal')

def button_generate_report_fun():
 
    button_choose_dir.config(state='disable')
    button_process.config(state='disable')
    t1= Thread(target=thread_fun )        
    t1.start()
    
    

    # messagebox.showinfo("Notice", "Report was generated")
def create_window():

    root = Tk()
    #root.withdraw()
    root.resizable(1,1)
    root.title("ReportGenerator")
    root.geometry('800x350+700+200')
    frame1 =Frame(root)
    global button_choose_dir
    button_choose_dir = Button(frame1, text="Choose Directory", command=button_choose_dir_fun) 
    button_choose_dir.grid() 

    #label_dir= Label(root,text='Please choose test data folder')
    #label_dir.grid(column=1, row=0)

    scrollbar_h = Scrollbar(root,orient=HORIZONTAL)
    #Label(root,text='Test Case List').grid(column=0,row=1)
    listbox_test_case = Listbox(root, bg="white", fg="black", bd=5, height=10, width=20, font=("Arial", 14), selectmode='multiple',xscrollcommand=scrollbar_h)
    listbox_test_case.pack(fill=X,padx=5)#grid(column=1, row=2)
    scrollbar_h.config(command=listbox_test_case.xview)
    scrollbar_h.pack(fill=X)
    global button_process
    button_process=Button(frame1,text='Generate Report',command=button_generate_report_fun)
    button_process.grid(column=1,row=0,padx=20)
    button_exit=Button(frame1,text='Exit',command=sys.exit)    
    button_exit.grid(column=2,row=0,padx=10)

    frame1.pack(pady=20)  


    global textvr_global   
    textvr_global= StringVar()
    label_status = Label(root,textvariable= textvr_global)
    label_status.pack()

    global listbox_global
    listbox_global = listbox_test_case

    root.mainloop() 
if __name__=="__main__":
    
    total_doc_content = ReportContent('')  
    
    create_window()

    